import re
from pathlib import Path
from docx import Document


def extrair_dados(log):
    logs = []

    expressao = r"(\d{4}-\d{2}-\d{2})\s*(\d{2}:\d{2}:\d{2})\s*(\w+)\s*(.+)"

    resultados = re.findall(expressao, log)

    for resultado in resultados:
        logs.append({
            'data': resultado[0],
            'hora': resultado[1],
            'tipo': resultado[2],
            'mensagem': resultado[3]
        })

    return logs


def criar_relatorio():

    # Pega a pasta onde o main.py está localizado
    pasta = Path(__file__).parent

    # Caminho do arquivo de logs
    arquivo_logs = pasta / "logs.txt"

    # Caminho do relatório
    arquivo_relatorio = pasta / "relatorio_teste.docx"

    # Verifica se logs.txt existe
    if not arquivo_logs.exists():
        print(f"Erro: o arquivo {arquivo_logs} não foi encontrado.")
        return

    # Lê o arquivo de logs
    with open(arquivo_logs, encoding='utf-8') as log:
        log_completo = log.read()

    # Extrai os dados
    logs = extrair_dados(log_completo)

    if not logs:
        print("Não encontrei nenhum log válido!")
        return

    # Cria o documento Word
    documento = Document()

    documento.add_heading(
        'Relatório de Análise de Logs',
        level=1
    )

    # Conta quantidade de cada tipo
    logs_tipos_qtd = {}

    for log in logs:

        tipo = log['tipo']

        if tipo in logs_tipos_qtd:
            logs_tipos_qtd[tipo] += 1
        else:
            logs_tipos_qtd[tipo] = 1

    documento.add_paragraph(
        "Total de ocorrências para cada tipo:"
    )

    for tipo, qtd in logs_tipos_qtd.items():
        documento.add_paragraph(
            f"{tipo}: {qtd}"
        )

    # Lista os erros
    documento.add_paragraph("Erros ocorridos:")

    encontrou_erro = False

    for log in logs:

        if log['tipo'] == "ERROR":

            encontrou_erro = True

            documento.add_paragraph(
                f"{log['data']} {log['hora']}: "
                f"{log['mensagem']}"
            )

    if not encontrou_erro:
        documento.add_paragraph(
            "Nenhum erro encontrado."
        )

    documento.add_page_break()

    # Organiza registros por dia
    registros_por_dia = {}

    for log in logs:

        data = log['data']
        tipo = log['tipo']

        if data not in registros_por_dia:
            registros_por_dia[data] = {}

        if tipo not in registros_por_dia[data]:
            registros_por_dia[data][tipo] = 0

        registros_por_dia[data][tipo] += 1

    # Cria tabela
    tabela = documento.add_table(
        rows=0,
        cols=5
    )

    tabela.style = 'Light List Accent 1'

    # Cabeçalho
    linha = tabela.add_row().cells

    linha[0].text = 'Data'
    linha[1].text = 'INFO'
    linha[2].text = 'ERROR'
    linha[3].text = 'WARNING'
    linha[4].text = 'DEBUG'

    # Preenche tabela
    for data in sorted(registros_por_dia):

        linha = tabela.add_row().cells

        linha[0].text = data
        linha[1].text = str(
            registros_por_dia[data].get('INFO', 0)
        )
        linha[2].text = str(
            registros_por_dia[data].get('ERROR', 0)
        )
        linha[3].text = str(
            registros_por_dia[data].get('WARNING', 0)
        )
        linha[4].text = str(
            registros_por_dia[data].get('DEBUG', 0)
        )

    # Salva o Word
    documento.save(arquivo_relatorio)

    print("Relatório criado com sucesso!")
    print(f"Arquivo: {arquivo_relatorio}")


criar_relatorio()